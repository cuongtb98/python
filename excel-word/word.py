from docx import Document


class MeetingPaper():
    def __init__(self,name, grade, rating):
        self.name = name
        self.grade = grade
        self.rating = rating
          

    def content(self):
        document = Document()
        document.add_heading('GIẤY MỜI HỌP PHỤ HUYNH', 0).alignment  = 1

        student = document.add_paragraph("Học sinh: ")
        student.add_run(self.name).bold = True
        student.add_run(" - Lớp: {}".format(self.grade))
        document.add_paragraph("Xếp loại học tập: {}".format(self.rating))

        time = document.add_paragraph("Thời gian: ")
        time.add_run("8AM ngày 04/11/2020")

        add = document.add_paragraph("Địa điểm: ")
        add.add_run("Phòng 303 trường THPT nam tiền hải")

        document.add_heading("Nội dung", 1)
        document.add_paragraph("Tổng kết năm học 2021 vừa qua và nhận xét quá trình học tập cửa từng học sinh")

        document.add_paragraph("Lộ trình năm học 2020", style="List Bullet")
        document.add_paragraph("Trao thưởng học sinh có thành tích học tập tốt", style="List Bullet")
        document.add_paragraph("Giải thích các khoản thu", style="List Bullet")

        document.add_heading("Các khoản thu", level=1)
        document.add_paragraph("Học phí: 3 triệu VNĐ", style="List Number")
        document.add_paragraph("Đồng phục: 1 triệu VNĐ", style="List Number")
        document.add_paragraph("Dã ngoại: 1 triệu VNĐ", style="List Number")
        document.add_heading("Tổng thu: 5 triệu VNĐ", level=2)

        document.add_paragraph("")
        document.add_paragraph("Giáo viên chủ nhiệm")
        document.add_paragraph("   Nguyễn Duy Long  ")
        document.save("giay_moi_hop_phu_huynh.docx")